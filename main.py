from docxtpl import DocxTemplate
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer
import os
import re


'''### TEST
test_template_path = "templates/testtest.docx"
doc = DocxTemplate(test_template_path)

context = {"name":"Oleksii"}

doc.render(context)
output_path = "output/test_output.docx"
os.makedirs("output", exist_ok=True)
doc.save(output_path)

print(f"Generated file saved at: {output_path}")
'''


# Constants
TEMPLATE_DIR = "templates"
OUTPUT_DIR = "output"

#Checking if we have blank pages
def is_blank_docx(doc_path):
    doc = Document(doc_path)
    return all(p.text.strip() == "" for p in doc.paragraphs)

#Taking Input for different templates
def get_user_input(template_type):
    data = {}
    data['EventName'] = input("Enter Event Name: ")
    data['Dates'] = input("Enter Dates (e.g. 07/24/2025): ")
    data['CellPhone'] = input("Enter Cell Phone #: ")

    if template_type == "boh":
        pass_count = int(input("How many passes? "))
        return data, pass_count
    elif template_type == "foh":
        pass_count = int(input("How many passes? "))
        return data, pass_count
    else:  # load/unload
        data['TimeEntered'] = input("Enter BOH Time: ")
        data['Deck'] = input("Enter Deck Number: ")
        pass_count = int(input("How many passes? "))
        return data, pass_count

#Creating a final document
def create_output_doc(template_path, data, pass_count=1):
    temp_files = []

    for i in range(1, pass_count + 1):
        data['PassNumber'] = i
        doc = DocxTemplate(template_path)  # ← moved inside loop
        doc.render(data)

        temp_path = os.path.join(OUTPUT_DIR, f"temp_{i}.docx")
        doc.save(temp_path)
        temp_files.append(temp_path)

    return temp_files


def combine_docs(temp_files, final_path):
    from docx import Document

    non_blank_files = [f for f in temp_files if not is_blank_docx(f)]
    if not non_blank_files:
        print("❗ No non-blank pages found.")
        return

    master = Document(non_blank_files[0])
    composer = Composer(master)

    for temp_file in non_blank_files[1:]:
        composer.append(Document(temp_file))

    composer.save(final_path)

    # Clean up blank temp files
    blank_files = [f for f in temp_files if is_blank_docx(f)]
    for f in blank_files:
        os.remove(f)



#Main Logic
def main():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    print("Select Template:")
    print("1 - FOH Pass")
    print("2 - BOH Pass")
    print("3 - Load/Unload Pass")
    choice = input("Enter 1, 2 or 3: ").strip()

    if choice == "1":
        template_file = "foh_pass_template.docx"
        template_type = "foh"
    elif choice == "2":
        template_file = "boh_pass_template.docx"
        template_type = "boh"
    elif choice == "3":
        template_file = "load_unload_template.docx"
        template_type = "load"
    else:
        print("Invalid choice.")
        return

    template_path = os.path.join(TEMPLATE_DIR, template_file)
    data, pass_count = get_user_input(template_type)

    temp_docs = create_output_doc(template_path, data, pass_count)

    safe_name = re.sub(r'[\\/*?:"<>|]', "_", data['EventName'])
    output_file = os.path.join(OUTPUT_DIR, f"{safe_name}_Passes.docx")

    if len(temp_docs) == 1:
        os.rename(temp_docs[0], output_file)
    else:
        combine_docs(temp_docs, output_file)
        for f in temp_docs:
            os.remove(f)

    print(f"✅ Done! File saved to: {output_file}")

if __name__ == "__main__":
    main()
