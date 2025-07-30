import streamlit as st
from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
import os
import re
import tempfile
import pypandoc


TEMPLATE_DIR = "templates"

def sanitize_filename(name):
    return re.sub(r'[\\/*?:"<>|]', "_", name)

def is_blank_docx(doc_path):
    doc = Document(doc_path)
    return all(p.text.strip() == "" for p in doc.paragraphs)

def create_single_doc(template_path, context, count):
    temp_files = []
    for i in range(1, count + 1):
        context['PassNumber'] = i
        doc = DocxTemplate(template_path)
        doc.render(context)

        path = os.path.join(tempfile.gettempdir(), f"temp_{i}.docx")
        doc.save(path)
        temp_files.append(path)
    return temp_files

def combine_docs(temp_files, final_path):
    master = Document(temp_files[0])
    composer = Composer(master)
    for file in temp_files[1:]:
        composer.append(Document(file))
    composer.save(final_path)

def convert_to_pdf(input_docx, output_pdf):
    try:
        pypandoc.convert_file(input_docx, 'pdf', outputfile=output_pdf)
        return output_pdf
    except OSError:
        st.error("❌ PDF conversion failed. Pandoc is not installed in this environment.")
        return None

def main():
    st.title("🎫 Parking Pass Generator")

    # Initialize session state for login
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if not st.session_state.authenticated:
        password = st.text_input("Enter Access Key", type="password")
        if password == "GSDPPass2025!":
            st.session_state.authenticated = True
            st.rerun()  # Refresh to hide the password field
        elif password:
            st.warning("Access denied. Please enter the correct access key.")
        return  # Exit app until access is granted

    template_type = st.selectbox("Select Pass Type", ["FOH Pass", "BOH Pass", "Load/Unload Pass"])
    event_name = st.text_input("Event Name")
    dates = st.text_input("Dates (e.g. 07/24/2025)")
    phone = st.text_input("Cell Phone #")
    num_passes = st.number_input("Number of Passes", min_value=1, value=1)

    additional_fields = {}
    if template_type == "Load/Unload Pass":
        additional_fields['TimeEntered'] = st.text_input("BOH Entry Time")
        additional_fields['Deck'] = st.text_input("Deck")

    export_format = st.radio("Export Format", ["DOCX"])

    if st.button("Generate"):
        with st.spinner("Creating passes..."):
            template_map = {
                "FOH Pass": "foh_pass_template.docx",
                "BOH Pass": "boh_pass_template.docx",
                "Load/Unload Pass": "load_unload_template.docx",
            }
            template_file = os.path.join(TEMPLATE_DIR, template_map[template_type])
            context = {
                "EventName": event_name,
                "Dates": dates,
                "CellPhone": phone,
                **additional_fields
            }

            temp_files = create_single_doc(template_file, context, num_passes)
            final_name = sanitize_filename(f"{event_name}_Passes")
            docx_path = os.path.join(tempfile.gettempdir(), f"{final_name}.docx")
            combine_docs(temp_files, docx_path)

            # Export logic
            if export_format == "PDF":
                output_pdf = docx_path.replace(".docx", ".pdf")
                result = convert_to_pdf(docx_path, output_pdf)
                if result:
                    with open(output_pdf, "rb") as f:
                        st.download_button("📄 Download PDF", f, file_name=os.path.basename(output_pdf))
                else:
                    st.warning("⚠️ PDF conversion failed. Serving DOCX instead.")
                    export_format = "DOCX"

            if export_format == "DOCX":
                with open(docx_path, "rb") as f:
                    st.download_button("📄 Download DOCX", f, file_name=os.path.basename(docx_path))

            # Cleanup
            for f in temp_files:
                os.remove(f)

if __name__ == "__main__":
    main()
